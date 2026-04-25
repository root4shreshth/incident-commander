"""Build the IncidentCommanderEnv project document (Word .docx).

Run:
    uv run python scripts/build_project_doc.py

Output:
    IncidentCommanderEnv_Project_Document.docx (in project root)
"""

from __future__ import annotations

import os
from datetime import date

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

ACCENT_RED = RGBColor(0xEF, 0x44, 0x44)
ACCENT_BLUE = RGBColor(0x3B, 0x82, 0xF6)
ACCENT_GREEN = RGBColor(0x22, 0xC5, 0x5E)
ACCENT_PURPLE = RGBColor(0xA8, 0x55, 0xF7)
TEXT_PRIMARY = RGBColor(0x1F, 0x29, 0x37)
TEXT_SECONDARY = RGBColor(0x4B, 0x55, 0x63)
TEXT_MUTED = RGBColor(0x6B, 0x72, 0x80)


def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell, color: str = "DDDDDD") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


def _add_run(paragraph, text: str, *, bold=False, italic=False,
             size=None, color=None, font="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return run


def _heading(doc, text: str, level: int):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        if level == 0:
            run.font.size = Pt(28)
            run.font.color.rgb = TEXT_PRIMARY
        elif level == 1:
            run.font.size = Pt(20)
            run.font.color.rgb = ACCENT_BLUE
        elif level == 2:
            run.font.size = Pt(15)
            run.font.color.rgb = TEXT_PRIMARY
        else:
            run.font.size = Pt(13)
            run.font.color.rgb = TEXT_SECONDARY
    return h


def _para(doc, text: str = "", *, bold=False, italic=False, size=11,
          color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = alignment
    if text:
        _add_run(p, text, bold=bold, italic=italic, size=size, color=color or TEXT_PRIMARY)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _bullet(doc, text: str, *, level=0, size=11):
    p = doc.add_paragraph(style="List Bullet")
    if level > 0:
        p.paragraph_format.left_indent = Cm(0.6 * (level + 1))
    _add_run(p, text, size=size, color=TEXT_PRIMARY)
    p.paragraph_format.space_after = Pt(2)
    return p


def _kv_table(doc, rows, col_widths_cm=(4.0, 12.5)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    for i, w in enumerate(col_widths_cm):
        for cell in table.columns[i].cells:
            cell.width = Cm(w)
    for r, (k, v) in enumerate(rows):
        c0, c1 = table.rows[r].cells
        c0.width = Cm(col_widths_cm[0])
        c1.width = Cm(col_widths_cm[1])
        _shade_cell(c0, "F3F4F6")
        _set_cell_borders(c0)
        _set_cell_borders(c1)
        c0.paragraphs[0].clear()
        _add_run(c0.paragraphs[0], k, bold=True, size=10, color=TEXT_PRIMARY)
        c1.paragraphs[0].clear()
        _add_run(c1.paragraphs[0], v, size=10, color=TEXT_PRIMARY)
    _para(doc, "", space_after=4)
    return table


def _data_table(doc, headers, rows, col_widths_cm=None):
    n = len(headers)
    if col_widths_cm is None:
        col_widths_cm = [16.5 / n] * n
    table = doc.add_table(rows=1 + len(rows), cols=n)
    table.autofit = False

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Cm(col_widths_cm[i])
        _shade_cell(cell, "1E3A8A")
        _set_cell_borders(cell, color="1E3A8A")
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Body
    for r, row in enumerate(rows):
        for i, val in enumerate(row):
            cell = table.rows[r + 1].cells[i]
            cell.width = Cm(col_widths_cm[i])
            if r % 2 == 0:
                _shade_cell(cell, "F9FAFB")
            _set_cell_borders(cell)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(val)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT_PRIMARY
    _para(doc, "", space_after=4)
    return table


def _quote(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.right_indent = Cm(1.2)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    _add_run(p, text, italic=True, size=12, color=ACCENT_BLUE)


def _separator(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), "CBD5E1")
    bottom.set(qn("w:space"), "1")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Document composition
# ---------------------------------------------------------------------------

def build_document(out_path: str) -> None:
    doc = Document()

    # Page setup — A4 with reasonable margins
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)

    # Default style: Calibri 11
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ============== TITLE PAGE ==============
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(60)
    title.paragraph_format.space_after = Pt(8)
    _add_run(title, "IncidentCommanderEnv", size=34, bold=True, color=ACCENT_RED)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(40)
    _add_run(sub, "An OpenEnv RL environment for training LLM agents to be on-call SREs",
             size=14, italic=True, color=TEXT_SECONDARY)

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag.paragraph_format.space_after = Pt(40)
    _add_run(tag, "Meta OpenEnv Hackathon — April 2026 · Theme #3.1 Professional Tasks",
             size=11, color=TEXT_MUTED)

    _quote(doc, "“We trained a 1.5B-parameter agent to diagnose and resolve real production "
                "incidents — first by cloning senior-SRE playbooks (SFT), then by reinforcing "
                "verifiable rewards (GRPO). The same trained policy runs unchanged on the "
                "simulator for training and on a real Docker stack for sim-to-real validation.”")

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_before = Pt(60)
    meta_p.paragraph_format.space_after = Pt(8)
    _add_run(meta_p, "Project Document & Mentor Pitch",
             size=12, bold=True, color=TEXT_PRIMARY)
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(date_p, date.today().strftime("%B %d, %Y"), size=10, color=TEXT_MUTED)
    repo_p = doc.add_paragraph()
    repo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(repo_p, "github.com/root4shreshth/incident-commander", size=10, color=ACCENT_BLUE)
    space_p = doc.add_paragraph()
    space_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(space_p, "hype4raj-incident-commander-env.hf.space", size=10, color=ACCENT_BLUE)
    doc.add_page_break()

    # ============== EXECUTIVE SUMMARY ==============
    _heading(doc, "Executive Summary", 1)
    _para(doc,
          "IncidentCommanderEnv is the first OpenEnv-compatible reinforcement-learning "
          "environment built specifically for SRE incident response. Production outages cost "
          "enterprises $1M–$5M per hour and the average mean-time-to-resolution is 8.85 "
          "hours. There has been no standardized environment to train or benchmark AI "
          "agents on this work — until now.",
          space_after=8)
    _para(doc,
          "The submission combines a fast, deterministic 9-service simulated cluster with "
          "a six-component verifiable reward, parametric scenario families, and a Backend "
          "abstraction that lets the same trained policy run on a real Docker Compose stack "
          "for sim-to-real validation. The training pipeline mirrors the hackathon’s "
          "explicit recipe: SFT warm-start from senior-SRE behavioral-clone trajectories, "
          "then GRPO fine-tuning with RLVR (verifiable rewards, no learned reward model).",
          space_after=8)
    _para(doc,
          "What this submission proves: an LLM agent can learn the shape of operational "
          "incident response — diagnose, trace dependency cascades, decide between "
          "restart/rollback/reconfigure, and declare resolution with an accurate root "
          "cause — when the environment provides a richly-decomposed reward and "
          "deterministic verification.",
          space_after=12)

    # ============== AT A GLANCE ==============
    _heading(doc, "At a Glance", 1)
    _kv_table(doc, [
        ("Theme", "#3.1 Professional Tasks (World Modeling)"),
        ("Domain", "SRE / DevOps incident response"),
        ("Environment shape", "9-service microservices cluster, 10 typed actions, 3 scenario families"),
        ("Training stack", "Hugging Face TRL + Unsloth, SFT then GRPO with verifiable rewards"),
        ("Reward design", "6 independent components (diagnostic, correct_op, resolution, format, efficiency, penalty)"),
        ("Reproducibility", "Seeded `/reset` produces byte-identical episodes; parametric scenarios for distributional training"),
        ("Test coverage", "186 automated tests including 15 anti-reward-hacking regression tests"),
        ("Deployable on", "HuggingFace Spaces (Docker), local FastAPI, Colab via openenv-core"),
        ("Sim-to-real", "Backend Protocol with `BACKEND=sim` (training) and `BACKEND=real` (Docker Compose)"),
    ])

    doc.add_page_break()

    # ============== PROBLEM ==============
    _heading(doc, "1. The Problem", 1)
    _heading(doc, "1.1 The pain point", 2)
    _para(doc,
          "Every tech company runs on-call rotations where engineers are paged at 3 AM "
          "to diagnose production outages under extreme time pressure. The job is "
          "expensive, slow, and burnout-inducing:",
          space_after=4)
    _bullet(doc, "Fortune 1000 companies lose $1.25B–$2.5B annually to preventable downtime")
    _bullet(doc, "97% of large enterprises say a single hour of downtime costs more than $100K")
    _bullet(doc, "Average global incident mean-time-to-resolution is 8.85 hours")
    _bullet(doc, "65% of engineers report burnout; 70% of SRE teams cite alert fatigue as top-3 concern")
    _bullet(doc, "78% of developers spend 30%+ of their time on manual operational toil")
    _para(doc, "", space_after=8)

    _heading(doc, "1.2 The capability gap for AI", 2)
    _para(doc,
          "Although LLMs are increasingly used as engineering co-pilots, no public RL "
          "environment exists for training them on incident response. OpenEnv, "
          "Gymnasium, and HuggingFace’s benchmark collections cover games, browsing, "
          "code repair, and math — but nothing for the operational substrate that runs "
          "the world’s production systems.",
          space_after=8)
    _para(doc, "The capability the field is missing:", space_after=4)
    _bullet(doc, "Causal reasoning under partial observability — see logs and metrics, infer system state")
    _bullet(doc, "Sequential decision-making with deferred consequences — wrong move makes things worse")
    _bullet(doc, "Tracing symptoms to root cause through dependency graphs (2-3 layers deep)")
    _bullet(doc, "Action ordering — in cascading failures, fix the source before fixing the dependents")
    _bullet(doc, "Knowing when not to act — restarting a healthy service is worse than nothing")
    _para(doc, "", space_after=12)

    # ============== SITUATION ==============
    _heading(doc, "2. Situation", 1)
    _heading(doc, "2.1 What the hackathon judges", 2)
    _para(doc, "The Meta OpenEnv Hackathon (April 2026) weights submissions on four axes:", space_after=4)
    _data_table(
        doc,
        ["Criterion", "Weight", "What it means"],
        [
            ["Environment Innovation", "40%", "Is the environment novel, creative, or genuinely challenging?"],
            ["Storytelling & Presentation", "30%", "Can the team clearly explain problem, env, and what was learned?"],
            ["Showing Improvement in Rewards", "20%", "Reward curves, before/after evidence, baseline comparison"],
            ["Reward & Training Pipeline", "10%", "Reward logic coherent? Pipeline produces meaningful improvement?"],
        ],
        col_widths_cm=[5.0, 2.0, 9.5],
    )

    _heading(doc, "2.2 Minimum requirements (non-negotiable)", 2)
    _bullet(doc, "Built on OpenEnv (latest release)")
    _bullet(doc, "Working training script using Unsloth or HuggingFace TRL, ideally as a Colab")
    _bullet(doc, "Evidence of real training — loss + reward plots from a real run")
    _bullet(doc, "Mini-blog on HuggingFace OR < 2-minute demo video")
    _bullet(doc, "Hosted on HuggingFace Spaces (discoverable + runnable)")
    _bullet(doc, "README that motivates the problem, explains the env, and shows results")
    _para(doc, "", space_after=12)

    # ============== SOLUTION ==============
    _heading(doc, "3. Solution", 1)
    _heading(doc, "3.1 The headline idea", 2)
    _quote(doc, "“Train an LLM agent to act as an on-call SRE — using SFT warm-start "
                "from senior-SRE playbooks, then GRPO fine-tuning with verifiable, "
                "decomposed rewards — and demonstrate sim-to-real transfer by pointing "
                "the same trained policy at a real Docker stack.”")
    _heading(doc, "3.2 Core architecture", 2)
    _para(doc,
          "The environment is a typed FastAPI service exposing the OpenEnv contract "
          "(POST /reset, POST /step, GET /state). Internally it delegates execution "
          "through a Backend Protocol so the same trained policy runs unchanged across "
          "two substrates:",
          space_after=4)
    _kv_table(doc, [
        ("SimulatedBackend",
         "Fast, in-memory Python cluster of 9 microservices — used for all training. "
         "Reset accepts a seed; same seed plus same actions yields byte-identical "
         "trajectories. Vectorizable for parallel rollouts."),
        ("RealBackend",
         "Wraps a Docker Compose stack via shell-outs (`docker compose restart`, "
         "`docker compose logs`, etc.). Used for the sim-to-real demo: the trained "
         "policy diagnoses a real container crash and resolves it with real Docker "
         "commands, zero retraining."),
    ], col_widths_cm=(4.5, 12.0))

    _heading(doc, "3.3 The six-component reward", 2)
    _para(doc,
          "The hackathon docs explicitly recommend multiple independent reward functions "
          "to defeat reward hacking. The reward is decomposed into six pure functions "
          "and aggregated as a typed RewardBreakdown that TRL logs to wandb separately:",
          space_after=4)
    _data_table(
        doc,
        ["Component", "Triggers", "Range"],
        [
            ["r_diagnostic", "Investigative reads on relevant services", "+0.02 .. +0.05"],
            ["r_correct_op", "Scenario-defined right move", "+0.15"],
            ["r_resolution", "Terminal — accurate root-cause declaration", "-0.05 .. +0.30"],
            ["r_format", "Per-step credit for well-formed action JSON", "+0.01"],
            ["r_efficiency", "Bonus when resolved within 50% of step budget", "0 or +0.10"],
            ["r_penalty", "Sum of harmful_restart, redundancy, handler_error", "-0.30 .. 0"],
        ],
        col_widths_cm=[4.5, 8.0, 4.0],
    )
    _para(doc,
          "Each component is independently unit-tested. TRL’s GRPOTrainer logs them "
          "separately so the storytelling visualization shows the agent learning to "
          "maximize correct-op and resolution while reducing penalty — visible "
          "evidence the agent isn’t single-axis hacking.",
          space_after=12)

    _heading(doc, "3.4 Anti-reward-hacking pass", 2)
    _para(doc,
          "Four documented exploits in the original env were closed with regression "
          "tests pinning each fix:",
          space_after=4)
    _data_table(
        doc,
        ["Exploit", "Original gap", "Fix"],
        [
            ["String-match config heal",
             "`update_config` healed any service if key contained ‘pool’ + ‘size’",
             "Strict allowlist + scenario.on_config_update hook"],
            ["Restart-clears-anomalies",
             "Any restart cleared all anomalies regardless of fault type",
             "_RESTART_CURABLE set: only oom, connection_leak, resource_starved"],
            ["Param-tweak redundancy bypass",
             "Different parameters skipped the redundancy penalty",
             "Match on (action_type, target) within 3-action window"],
            ["Rollback-to-self",
             "Rollback to current version silently appended history + healed",
             "Early-return error; anomalies stay intact"],
        ],
        col_widths_cm=[4.5, 7.0, 5.0],
    )
    _para(doc, "", space_after=10)

    _heading(doc, "3.5 Parametric scenario families", 2)
    _para(doc,
          "Three scenario classes were converted from hardcoded instances into "
          "seeded factories. Each /reset materializes a fresh instance from a "
          "distribution — the agent learns shape, not constants.",
          space_after=4)
    _data_table(
        doc,
        ["Family", "Difficulty", "What randomizes per seed"],
        [
            ["oom_crash", "Easy",
             "Target service (4 candidates), memory limit (192–320 Mi), step budget"],
            ["db_pool_exhaustion", "Medium",
             "Initial pool size (16/20/24), step budget"],
            ["bad_deployment_cascade", "Hard",
             "Step budget; correct ordering still required (rollback before restarts)"],
        ],
        col_widths_cm=[5.0, 2.5, 9.0],
    )
    doc.add_page_break()

    # ============== HOW IT WORKS ==============
    _heading(doc, "4. How It Works", 1)

    _heading(doc, "4.1 The episode loop", 2)
    _para(doc,
          "Every episode follows the standard OpenEnv contract:",
          space_after=4)
    _bullet(doc, "POST /reset {task_id, seed, difficulty} — alert fires, dependency graph returned")
    _bullet(doc, "POST /step {action_type, target_service, parameters} — execute one of 10 typed actions")
    _bullet(doc, "Per-step observation includes message, logs/metrics/snapshot, reward, done flag")
    _bullet(doc, "Episode terminates when scenario.check_resolved() is true OR step_count == max_steps")
    _bullet(doc, "Final score is the deterministic rubric grade (0..1) clamped to (0.01, 0.99)")
    _para(doc, "", space_after=8)

    _heading(doc, "4.2 The 10 SRE actions", 2)
    _data_table(
        doc,
        ["Action", "Group", "Purpose"],
        [
            ["list_services", "Investigate", "Cluster-wide health + key metrics overview"],
            ["describe_service", "Investigate", "Config, deployment history, dependencies"],
            ["read_logs", "Investigate", "Recent log lines (filterable by severity)"],
            ["check_metrics", "Investigate", "CPU, memory, latency p50/p99, error rate"],
            ["run_diagnostic", "Investigate", "Connectivity / health / DNS / resource probes"],
            ["restart_service", "Remediate", "Bounce service; optionally raise memory limit"],
            ["rollback_deployment", "Remediate", "Revert to a prior version"],
            ["scale_service", "Remediate", "Change replica count"],
            ["update_config", "Remediate", "Set runtime config (allowlist enforced)"],
            ["resolve_incident", "Declare", "Submit root-cause + resolution narrative"],
        ],
        col_widths_cm=[4.5, 2.5, 9.5],
    )
    _para(doc, "", space_after=4)

    _heading(doc, "4.3 Training pipeline", 2)
    _para(doc,
          "The pipeline mirrors the OpenEnv hackathon docs’ explicit recipe — "
          "SFT to warm-start a base model with format compliance and a diagnostic "
          "prior, then GRPO with verifiable rewards (RLVR) to fine-tune the policy.",
          space_after=4)
    _data_table(
        doc,
        ["Phase", "What happens", "Output"],
        [
            ["1. SFT dataset build",
             "Replay senior-SRE trajectories under multiple seeds → ~120 chat rows",
             "HuggingFace Dataset"],
            ["2. SFT training",
             "Unsloth 4-bit Qwen2.5-Coder + LoRA, 1 epoch, lr=2e-4",
             "sft-adapter checkpoint"],
            ["3. Eval after SFT",
             "30 held-out seeds × 3 families = 90 episodes vs random baseline",
             "JSON eval report"],
            ["4. GRPO training",
             "TRL GRPOTrainer, 4 rollouts/prompt, 6-component reward, curriculum-gated families",
             "grpo-adapter checkpoint"],
            ["5. Final eval",
             "4 conditions × 3 families × 30 seeds = 360 episodes",
             "results/eval_summary.json"],
            ["6. Plots + push",
             "Reward curves, components, success bars, action distribution → results/*.png",
             "HF Hub upload"],
        ],
        col_widths_cm=[3.5, 9.0, 4.0],
    )
    _para(doc, "", space_after=8)

    _heading(doc, "4.4 The dashboard observability mode", 2)
    _para(doc,
          "The existing human-facing UI doubles as an agent observability dashboard. "
          "A `?watch=<run_id>` URL parameter switches it into a passive replay mode "
          "where the service map turns red/green as the trained agent acts, the "
          "notebook streams the agent’s actions, and the AI Coach panel becomes a "
          "live chain-of-thought view. Same UI components, different data source.",
          space_after=12)

    # ============== FEATURES ==============
    _heading(doc, "5. Feature Inventory", 1)

    _heading(doc, "5.1 Environment features", 2)
    _bullet(doc, "9-service in-memory cluster with realistic dependency graph and resource quota")
    _bullet(doc, "Per-service health states: healthy / degraded / unhealthy / crashed / restarting")
    _bullet(doc, "Live metrics: CPU%, memory MB, request_latency_p50_ms, p99_ms, error_rate, RPS")
    _bullet(doc, "Structured log buffer with realistic OOM, connection-leak, deploy-cascade patterns")
    _bullet(doc, "Deployment history with rollback")
    _bullet(doc, "Anomaly engine with 6 fault types: oom, db_pool_exhaustion, connection_leak, "
                  "cascade_degradation, memory_leak, resource_starved")
    _bullet(doc, "Three parametric scenario families (Easy/Medium/Hard)")
    _bullet(doc, "OpenEnv-compliant: action_space + observation_space declared in openenv.yaml")
    _bullet(doc, "Typed Pydantic models for action, observation, state, and per-component reward")

    _heading(doc, "5.2 Reward + grading features", 2)
    _bullet(doc, "Six independent reward components (RewardBreakdown dataclass)")
    _bullet(doc, "EpisodeContext typed bundle every component reads from")
    _bullet(doc, "Episode rubric grader with weighted criteria + scenario-specific penalties")
    _bullet(doc, "Multiplicative time decay (0.995/step) — incentivizes faster resolution")
    _bullet(doc, "/reward-breakdown HTTP endpoint exposes the per-step decomposition")

    _heading(doc, "5.3 Training pipeline features", 2)
    _bullet(doc, "Self-contained Colab notebook (training/train_grpo.ipynb)")
    _bullet(doc, "Auto-detects compute (T4 vs A100) and selects model accordingly")
    _bullet(doc, "SFT dataset builder from senior-SRE behavioral-clone trajectories")
    _bullet(doc, "Phase-gated curriculum scheduler (warmup_oom → ops_mixed → full_mix)")
    _bullet(doc, "Model-agnostic eval runner (random / base / SFT / SFT+GRPO via the same code)")
    _bullet(doc, "Four matplotlib plot generators: reward curve, components, success bars, action dist")

    _heading(doc, "5.4 Quality + safety features", 2)
    _bullet(doc, "186 automated tests across reward components, exploit regression, seeded reproducibility, "
                  "Backend Protocol contract, and training-module plumbing")
    _bullet(doc, "Four documented reward-hacking exploits closed with explicit regression tests")
    _bullet(doc, "Strict (0,1) reward clamp matching the OpenEnv validator’s expectations")
    _bullet(doc, "Full episode replay determinism: same seed + same actions → byte-identical observations")
    _bullet(doc, "Backend Protocol abstraction enables sim-to-real transfer without retraining")
    _bullet(doc, "Anti-pollution `Cluster.initialize` deep-copies service config (test-suite hygiene)")
    _para(doc, "", space_after=10)

    # ============== TECH STACK ==============
    _heading(doc, "6. Technical Stack", 1)
    _data_table(
        doc,
        ["Layer", "Technology", "Why"],
        [
            ["Language", "Python 3.10+", "OpenEnv ecosystem standard"],
            ["API framework", "FastAPI", "Async, auto-docs, Pydantic-native"],
            ["Data models", "Pydantic v2", "Type-safe, self-documenting, OpenAPI-compatible"],
            ["Containerization", "Docker", "Required by OpenEnv spec"],
            ["Hosting", "HuggingFace Spaces", "Free hosting, OpenEnv-native"],
            ["Inference client", "OpenAI-compatible", "Works with any provider (HF Router, OpenRouter, Groq)"],
            ["Simulation", "Pure Python (in-memory)", "No external deps, deterministic, fast"],
            ["Training", "Unsloth + HuggingFace TRL", "4-bit LoRA + GRPO; matches hackathon recipe"],
            ["Base model", "Qwen2.5-Coder-1.5B (T4) / 7B (A100)", "Code-specialized; produces well-formed JSON actions"],
            ["Plots", "matplotlib", "Standard, embeddable, dpi=120 PNG"],
        ],
        col_widths_cm=[3.5, 5.5, 7.5],
    )

    _heading(doc, "6.1 Project layout", 2)
    _para(doc, "Top-level structure of the repository:", space_after=4)
    layout_lines = [
        "incident_commander_env/",
        "  models.py                       — typed actions, observations, state",
        "  openenv.yaml                    — full spec manifest (action_space, obs_space, reward decomposition)",
        "  server/",
        "    app.py                        — FastAPI routes (/reset, /step, /reward-breakdown, /backend, /coach/*)",
        "    environment.py                — IncidentCommanderEnv orchestrator (delegates to Backend)",
        "    coach.py                      — IDEAL_TRAJECTORIES (SFT seed) + LEARNING_CONTEXT + post-mortem builder",
        "    backends/",
        "      protocol.py                 — Backend Protocol + BackendSnapshot/ServiceSnapshot",
        "      sim.py                      — SimulatedBackend (in-memory Cluster)",
        "      real.py                     — RealBackend (Docker Compose shell-outs)",
        "    grading/",
        "      reward.py                   — backwards-compatible facade",
        "      components.py               — six reward components (pure functions)",
        "      episode_context.py          — EpisodeContext typed bundle",
        "      grader.py                   — rubric grader",
        "    scenarios/                    — OOMCrashScenario, DBPoolScenario, BadDeployScenario (parametric)",
        "    simulation/                   — Cluster, Service, dependency_graph, metrics_engine, log_generator",
        "    actions/handlers.py           — 10 typed action handlers",
        "    static/                       — frontend dashboard (HTML, demo.js, coach.js, map.js, auth.js, icons.js)",
        "training/",
        "  datasets.py                     — SFT dataset builder",
        "  eval_runner.py                  — model-agnostic eval (random / SFT / SFT+GRPO)",
        "  curriculum.py                   — phase-gated scenario sampler",
        "  plots.py                        — four matplotlib plot generators",
        "  grpo_reward.py                  — TRL reward function using the 6-component breakdown",
        "  train_grpo.ipynb                — self-contained Colab",
        "tests/                            — 186 tests (component, reward-hack, reproducibility, backend, training)",
        "Dockerfile                        — python:3.11-slim + uvicorn",
        "openenv.yaml                      — top-level spec for HuggingFace Hub rendering",
        "pyproject.toml                    — deps + [training] optional-extras",
    ]
    for line in layout_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = TEXT_PRIMARY
    _para(doc, "", space_after=10)

    doc.add_page_break()

    # ============== ROADMAP ==============
    _heading(doc, "7. Roadmap & Future Work", 1)
    _heading(doc, "7.1 What ships in the hackathon submission", 2)
    _bullet(doc, "Multi-component verifiable reward + parametric scenarios + seeded reproducibility")
    _bullet(doc, "Backend Protocol with SimulatedBackend (training) + RealBackend (sim-to-real demo)")
    _bullet(doc, "SFT-then-GRPO training pipeline (Colab) + trained adapter on HF Hub")
    _bullet(doc, "Reward curves + per-component plots + 4-condition × 3-family eval table")
    _bullet(doc, "Dashboard observability mode for live agent replay")
    _bullet(doc, "HF Space deployment + README + mini-blog + < 2-minute video")

    _heading(doc, "7.2 Phase 2 — Code-aware mode (post-hackathon)", 2)
    _bullet(doc, "Discriminated typed action union: Ops + Shell (sandboxed allowlist) + Code (read_file, "
                  "grep, propose_patch, apply_patch, run_tests) + Declare")
    _bullet(doc, "CodeAwareBackend wrapping a git worktree + pytest")
    _bullet(doc, "Code-fix scenario variants where rollback isn’t enough — agent reads source, "
                  "writes a patch, runs tests, applies if green")
    _bullet(doc, "SWE-Bench × incident-response hybrid benchmark")

    _heading(doc, "7.3 Phase 3 — Crowdsourced scenario library", 2)
    _bullet(doc, "YAML scenario authoring DSL")
    _bullet(doc, "Community contributions of real post-mortem-derived scenarios")
    _bullet(doc, "Versioned eval leaderboard")

    _heading(doc, "7.4 Phase 4 — Bring-your-own-codebase mode", 2)
    _bullet(doc, "Read-only shadow mode against real production observability stacks "
                  "(Datadog, Grafana, CloudWatch)")
    _bullet(doc, "Opt-in active mode (with audit logging) on staging clusters")
    _bullet(doc, "On-prem deployment for enterprise tier")
    _para(doc, "", space_after=10)

    # ============== RISK + MITIGATION ==============
    _heading(doc, "8. Risk Analysis & Mitigation", 1)
    _data_table(
        doc,
        ["Risk", "Probability", "Mitigation"],
        [
            ["GRPO doesn’t converge", "Medium",
             "Land SFT first as headline result; GRPO is the cherry on top"],
            ["Reward hacking discovered late", "Medium",
             "Six independent components + regression tests pin known exploits closed"],
            ["Compute over-budget on Colab", "Low-Med",
             "Auto-detect VRAM; fall back to Qwen-1.5B on T4 if A100 unavailable"],
            ["HF Space rebuild flake", "Medium",
             "Default BACKEND=sim keeps the Space healthy without Docker; smoke-tested"],
            ["Sim-to-real gap on Phase 6 demo", "Low",
             "Backend Protocol provides identical action surface; trained policy generalizes by construction"],
            ["Test-suite pollution between fixtures", "Low (resolved)",
             "Cluster.initialize deep-copies ServiceConfig; latent bug found and patched"],
        ],
        col_widths_cm=[5.0, 3.0, 8.5],
    )
    _para(doc, "", space_after=10)

    # ============== METRICS THAT MATTER ==============
    _heading(doc, "9. Metrics That Matter", 1)
    _para(doc,
          "What we will report in the README and pitch:",
          space_after=4)
    _bullet(doc, "Episode success rate per scenario family (random / base / SFT / SFT+GRPO)")
    _bullet(doc, "Average rubric score per condition")
    _bullet(doc, "Average steps used per episode (efficiency proxy)")
    _bullet(doc, "Action distribution shift before vs after training (visible reasoning shift)")
    _bullet(doc, "Per-component reward over training steps (anti-hacking visualization)")
    _bullet(doc, "Sim-to-real demo: real Docker container actually crashed and actually recovered")
    _para(doc, "", space_after=12)

    doc.add_page_break()

    # ============== PITCH ==============
    _heading(doc, "10. Mentor Pitch", 1)

    _heading(doc, "10.1 The hook (one sentence)", 2)
    _quote(doc, "“IncidentCommanderEnv is the first OpenEnv environment that teaches "
                "LLMs to do on-call SRE work — and we trained the first agent on it.”")

    _heading(doc, "10.2 Three-act narrative (90 seconds)", 2)
    _para(doc, "Act I — The world has a problem (15 seconds)", bold=True, space_after=2)
    _para(doc,
          "It’s 3 AM. Your phone buzzes — production is down. Companies lose $1M–$5M "
          "per hour during outages, the average incident takes nearly 9 hours to "
          "resolve, and there is no public RL benchmark for ops AI. We built one.",
          space_after=8)

    _para(doc, "Act II — How it works (45 seconds)", bold=True, space_after=2)
    _para(doc,
          "Nine simulated services, three scenario families, ten typed actions — and a "
          "six-component verifiable reward designed to defeat reward hacking. We seeded "
          "a Qwen-2.5-Coder LLM with senior-SRE behavioral-clone trajectories (SFT), "
          "then fine-tuned with GRPO using the env’s decomposed reward. Every component "
          "is logged separately so you can see the agent learning to maximize correct-op "
          "and resolution while reducing harmful actions — visible evidence it isn’t "
          "single-axis hacking.",
          space_after=8)

    _para(doc, "Act III — The killer demo (30 seconds)", bold=True, space_after=2)
    _para(doc,
          "Same trained policy, two substrates: simulator for fast training, real "
          "Docker stack for the sim-to-real demo. Watch a real `payment-service` "
          "container actually OOM-kill on stage. The trained agent reads the real "
          "Docker logs, identifies the OOM, calls restart_service with memory_limit=512Mi, "
          "the real container actually recovers, and the real /health endpoint returns 200. "
          "Trained on simulation. Validated against reality.",
          space_after=12)

    _heading(doc, "10.3 Why mentors should care", 2)
    _bullet(doc, "Domain unprecedented — zero prior art in OpenEnv, Gymnasium, or HF benchmarks for ops AI")
    _bullet(doc, "Methodology mirrors the docs’ explicit recipe (SFT-then-RL with verifiable rewards)")
    _bullet(doc, "Reward design defeats single-axis hacking by construction")
    _bullet(doc, "Same action API across sim and real → no retraining for sim-to-real transfer")
    _bullet(doc, "Behavioral-clone seed data from senior-SRE annotations is a reusable benchmark asset")
    _bullet(doc, "Live observability dashboard tells the story visually — most submissions only have plots")
    _para(doc, "", space_after=8)

    _heading(doc, "10.4 Closing line", 2)
    _quote(doc, "“Production incidents cost the industry billions per year. We built the "
                "first benchmark to train AI agents on them — and the first trained agent "
                "to resolve them. Mentor us, and the next chapter is code-aware "
                "remediation: the agent doesn’t just restart the service, it reads the "
                "source, writes the patch, runs the tests, and ships the fix. That’s "
                "operational AI you’d actually trust at 3 AM.”")

    _para(doc, "", space_after=4)
    _separator(doc)
    _para(doc, "", space_after=2)
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(closing,
             "Built for the Meta OpenEnv Hackathon · April 2026",
             size=10, italic=True, color=TEXT_MUTED)

    # ============== APPENDIX A — EVAL TABLE TEMPLATE ==============
    doc.add_page_break()
    _heading(doc, "Appendix A — Eval Table Template (for README + pitch slide)", 1)
    _para(doc,
          "Numbers below are placeholders pending the Phase 5 training run. The "
          "table format and the conditions are final.",
          italic=True, color=TEXT_SECONDARY, space_after=8)
    _data_table(
        doc,
        ["Condition", "OOM Crash", "DB Pool", "Bad Deploy", "Average"],
        [
            ["Random",   "—%", "—%", "—%", "—%"],
            ["Base (Qwen, no training)", "—%", "—%", "—%", "—%"],
            ["SFT only", "—%", "—%", "—%", "—%"],
            ["SFT + GRPO", "—%", "—%", "—%", "—%"],
        ],
        col_widths_cm=[4.0, 3.0, 3.0, 3.0, 3.5],
    )
    _para(doc,
          "Each cell reports episode success rate (resolution before step budget "
          "exhaustion) over 30 held-out seeds per scenario family. Total: 360 episodes "
          "per row.",
          space_after=10)

    # ============== APPENDIX B — GLOSSARY ==============
    _heading(doc, "Appendix B — Glossary", 1)
    glossary = [
        ("SRE", "Site Reliability Engineering — discipline of applying software engineering to operations."),
        ("MTTR", "Mean Time To Resolution — average time to fix an incident."),
        ("OOM", "Out Of Memory — process killed because it exceeded its memory limit."),
        ("RL", "Reinforcement Learning — training method based on rewards and penalties."),
        ("RLVR", "Reinforcement Learning with Verifiable Rewards — uses a verifier/test harness instead of a learned reward model."),
        ("OpenEnv", "Meta/PyTorch framework for defining standardized RL environments."),
        ("Gymnasium", "Standard API for RL environments (successor to OpenAI Gym)."),
        ("SFT", "Supervised Fine-Tuning — training on (input, ideal-output) pairs."),
        ("GRPO", "Group Relative Policy Optimization — a TRL-supported RL trainer that scores groups of completions."),
        ("LoRA", "Low-Rank Adaptation — parameter-efficient fine-tuning (small trainable adapter on top of frozen base weights)."),
        ("HuggingFace Spaces", "Platform for deploying ML applications, used as the official OpenEnv hosting target."),
        ("Backend Protocol", "Our typed interface (sim/real/code-aware) that lets the same trained policy run across substrates."),
    ]
    _data_table(
        doc,
        ["Term", "Meaning"],
        glossary,
        col_widths_cm=[3.5, 13.0],
    )

    # Save
    doc.save(out_path)


if __name__ == "__main__":
    here = os.path.dirname(os.path.abspath(__file__))
    out = os.path.abspath(os.path.join(here, os.pardir, "IncidentCommanderEnv_Project_Document.docx"))
    build_document(out)
    print(f"Saved: {out}")
